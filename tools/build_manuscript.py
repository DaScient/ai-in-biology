#!/usr/bin/env python3
"""Compile the AI in Biological Sciences textbook into a single styled ``.docx``.

This tool reads the chapter Markdown sources under ``docs/source`` in the order
declared by the MkDocs ``nav`` (preserving the "Part" groupings) and renders a
publication-ready Microsoft Word manuscript suitable for the DaScient, Inc.
textbook processing pipeline.

It preserves the structural elements of a textbook manuscript:

* Headings, subheadings and sections (mapped to Word heading styles).
* Equations and formulas (kept verbatim, including inline ``$...$`` math).
* Code blocks (rendered in a monospace, shaded style with a language label).
* Tables, ordered/unordered lists, block quotes and inline emphasis.

Project-specific front matter (ISBN, copyright, table of contents) is left as
clearly marked ``[TBD - DaScient, Inc.]`` placeholders so the publisher can drop
in the final values without touching the body content.

Usage
-----
    python tools/build_manuscript.py
    python tools/build_manuscript.py --output build/manuscript/AI_in_Biological_Sciences.docx

The script only depends on ``python-docx`` (already part of the docs toolchain
when installed) and ``PyYAML``.
"""

from __future__ import annotations

import argparse
import datetime as _dt
import re
from dataclasses import dataclass, field
from pathlib import Path

import yaml
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# --------------------------------------------------------------------------- #
# Paths and project metadata
# --------------------------------------------------------------------------- #

REPO_ROOT = Path(__file__).resolve().parents[1]
MKDOCS_YML = REPO_ROOT / "mkdocs.yml"
DOCS_DIR = REPO_ROOT / "docs" / "source"
DEFAULT_OUTPUT = REPO_ROOT / "manuscript" / "AI_in_Biological_Sciences.docx"

# Front-matter values the publisher (DaScient, Inc.) fills in later.
TBD = "[TBD - DaScient, Inc.]"

BOOK_TITLE = "AI in Biological Sciences"
BOOK_SUBTITLE = "Theory, Applications, Practice, and Society"
BOOK_AUTHORS = "Dr. Aris Thorne \u00b7 Wei Chen \u00b7 Marcus Adebayo"
PUBLISHER = "DaScient Press"

# Visual palette (matches the MkDocs Material teal/cyan theme).
TEAL = RGBColor(0x00, 0x69, 0x5C)
CODE_BG = "F2F4F3"
CODE_FONT = "Consolas"
BODY_FONT = "Calibri"


# --------------------------------------------------------------------------- #
# MkDocs navigation parsing
# --------------------------------------------------------------------------- #


@dataclass
class ChapterRef:
    """A single chapter, grouped under an optional Part heading."""

    part: str
    path: Path


def _load_mkdocs_nav() -> list[dict]:
    """Load the ``nav`` section from ``mkdocs.yml``.

    MkDocs uses ``!!python/name:`` tags that PyYAML's safe loader rejects, so we
    register harmless no-op constructors for the unknown tags.
    """

    class _MkDocsLoader(yaml.SafeLoader):
        pass

    def _ignore_unknown(loader, suffix, node):  # noqa: ANN001
        return None

    _MkDocsLoader.add_multi_constructor("tag:yaml.org,2002:python/name:", _ignore_unknown)
    _MkDocsLoader.add_multi_constructor("!!python/name:", _ignore_unknown)

    with MKDOCS_YML.open(encoding="utf-8") as fh:
        config = yaml.load(fh, Loader=_MkDocsLoader)
    return config.get("nav", [])


def collect_chapters() -> list[ChapterRef]:
    """Return the ordered list of chapters with their Part groupings."""
    nav = _load_mkdocs_nav()
    chapters: list[ChapterRef] = []

    for entry in nav:
        if not isinstance(entry, dict) or "Chapters" not in entry:
            continue
        for part_block in entry["Chapters"]:
            if not isinstance(part_block, dict):
                continue
            for part_name, items in part_block.items():
                for item in items:
                    rel = item if isinstance(item, str) else next(iter(item.values()))
                    chapters.append(ChapterRef(part=part_name, path=DOCS_DIR / rel))
    return chapters


# --------------------------------------------------------------------------- #
# Lightweight Markdown block model
# --------------------------------------------------------------------------- #


@dataclass
class Block:
    kind: str  # heading | paragraph | code | quote | list | table | blank
    text: str = ""
    level: int = 0
    language: str = ""
    items: list[tuple[int, str, str]] = field(default_factory=list)  # (indent, marker, text)
    rows: list[list[str]] = field(default_factory=list)


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_FENCE_RE = re.compile(r"^([ \t]*)(`{3,}|~{3,})\s*([\w+-]*)\s*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|?[\s:|-]+\|?\s*$")
_OL_RE = re.compile(r"^(\s*)(\d+)[.)]\s+(.*)$")
_UL_RE = re.compile(r"^(\s*)[-*+]\s+(.*)$")


def parse_markdown(md: str) -> list[Block]:
    """Parse a chapter's Markdown into a flat list of :class:`Block` objects."""
    lines = md.replace("\r\n", "\n").split("\n")
    blocks: list[Block] = []
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]

        # Fenced code block (``` or ~~~), preserving exact contents.
        fence = _FENCE_RE.match(line)
        if fence:
            fence_char = fence.group(2)[0]
            language = fence.group(3)
            i += 1
            body: list[str] = []
            while i < n and not re.match(rf"^[ \t]*{re.escape(fence_char)}{{3,}}\s*$", lines[i]):
                body.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            blocks.append(Block(kind="code", text="\n".join(body), language=language))
            continue

        # Headings.
        heading = _HEADING_RE.match(line)
        if heading:
            blocks.append(
                Block(kind="heading", level=len(heading.group(1)), text=heading.group(2).strip())
            )
            i += 1
            continue

        # Block quote (one or more consecutive ``>`` lines).
        if line.lstrip().startswith(">"):
            quote: list[str] = []
            while i < n and lines[i].lstrip().startswith(">"):
                quote.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            blocks.append(Block(kind="quote", text=" ".join(q.strip() for q in quote).strip()))
            continue

        # Tables (a header row followed by a separator row).
        if "|" in line and i + 1 < n and _TABLE_SEP_RE.match(lines[i + 1]):
            rows: list[list[str]] = [_split_table_row(line)]
            i += 2  # skip header + separator
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append(_split_table_row(lines[i]))
                i += 1
            blocks.append(Block(kind="table", rows=rows))
            continue

        # Lists (ordered or unordered, with simple nesting by indentation).
        if _OL_RE.match(line) or _UL_RE.match(line):
            items: list[tuple[int, str, str]] = []
            while i < n and (_OL_RE.match(lines[i]) or _UL_RE.match(lines[i])):
                ol = _OL_RE.match(lines[i])
                if ol:
                    indent = len(ol.group(1))
                    items.append((indent, "ol", ol.group(3).strip()))
                else:
                    ul = _UL_RE.match(lines[i])
                    indent = len(ul.group(1))
                    items.append((indent, "ul", ul.group(2).strip()))
                i += 1
            blocks.append(Block(kind="list", items=items))
            continue

        # Blank line.
        if not line.strip():
            blocks.append(Block(kind="blank"))
            i += 1
            continue

        # Paragraph: gather consecutive non-special lines.
        para: list[str] = []
        while i < n and lines[i].strip():
            nxt = lines[i]
            if (
                _HEADING_RE.match(nxt)
                or _FENCE_RE.match(nxt)
                or nxt.lstrip().startswith(">")
                or _OL_RE.match(nxt)
                or _UL_RE.match(nxt)
                or ("|" in nxt and i + 1 < n and _TABLE_SEP_RE.match(lines[i + 1]))
            ):
                break
            para.append(nxt.strip())
            i += 1
        blocks.append(Block(kind="paragraph", text=" ".join(para).strip()))

    return blocks


def _split_table_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


# --------------------------------------------------------------------------- #
# Inline Markdown rendering (bold / italic / code / links)
# --------------------------------------------------------------------------- #

_INLINE_RE = re.compile(
    r"(\*\*.+?\*\*|__.+?__|\*[^*]+?\*|_[^_]+?_|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))"
)


def add_inline(paragraph, text: str, *, base_italic: bool = False) -> None:
    """Append ``text`` to ``paragraph`` honouring inline Markdown emphasis."""
    pos = 0
    for match in _INLINE_RE.finditer(text):
        if match.start() > pos:
            _run(paragraph, text[pos : match.start()], italic=base_italic)
        token = match.group(0)
        if token.startswith("**") or token.startswith("__"):
            _run(paragraph, token[2:-2], bold=True, italic=base_italic)
        elif token.startswith("`"):
            _run(paragraph, token[1:-1], code=True)
        elif token.startswith("["):
            label = token[1 : token.index("]")]
            _run(paragraph, label, italic=base_italic)
        else:  # * or _ italic
            _run(paragraph, token[1:-1], italic=True)
        pos = match.end()
    if pos < len(text):
        _run(paragraph, text[pos:], italic=base_italic)


def _run(paragraph, text: str, *, bold=False, italic=False, code=False):
    if not text:
        return None
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if code:
        run.font.name = CODE_FONT
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
    return run


# --------------------------------------------------------------------------- #
# Word styling helpers
# --------------------------------------------------------------------------- #


def _shade_paragraph(paragraph, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    paragraph.paragraph_format.element.get_or_add_pPr().append(shd)


def _add_code_block(doc: Document, block: Block) -> None:
    label = "diagram" if block.language == "mermaid" else (block.language or "code")
    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(0)
    run = caption.add_run(f"\u27e6 {label} \u27e7")
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for raw in block.text.split("\n"):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.left_indent = Pt(12)
        _shade_paragraph(para, CODE_BG)
        run = para.add_run(raw if raw else "\u00a0")
        run.font.name = CODE_FONT
        run.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_table(doc: Document, block: Block) -> None:
    rows = block.rows
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            para = cell.paragraphs[0]
            add_inline(para, row[c_idx] if c_idx < len(row) else "")
            if r_idx == 0:
                for run in para.runs:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_list(doc: Document, block: Block) -> None:
    base_indent = min((indent for indent, _, _ in block.items), default=0)
    for indent, kind, text in block.items:
        level = 0 if indent <= base_indent else 1
        style = "List Number" if kind == "ol" else "List Bullet"
        if level:
            style += " 2"
        try:
            para = doc.add_paragraph(style=style)
        except KeyError:
            para = doc.add_paragraph(style="List Bullet")
        add_inline(para, text)


def _add_heading(doc: Document, block: Block) -> None:
    # Chapter title (``# ``) starts a new page. It keeps the built-in
    # "Heading 1" style (so it is picked up by Word's table-of-contents field)
    # while overriding the run formatting for the teal textbook look.
    if block.level == 1:
        doc.add_page_break()
        para = doc.add_heading(level=1)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(12)
        run = para.add_run(block.text)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = TEAL
        return
    heading = doc.add_heading(level=min(block.level, 4))
    add_inline(heading, block.text)


# --------------------------------------------------------------------------- #
# Front matter
# --------------------------------------------------------------------------- #


def _centered(doc: Document, text: str, *, size: int, bold=False, color=None, space_after=6):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return para


def add_front_matter(doc: Document) -> None:
    """Title page, copyright page and table-of-contents placeholders."""
    # --- Title page ---
    for _ in range(4):
        doc.add_paragraph()
    _centered(doc, BOOK_TITLE, size=32, bold=True, color=TEAL, space_after=4)
    _centered(doc, BOOK_SUBTITLE, size=16, space_after=24)
    _centered(doc, BOOK_AUTHORS, size=13, space_after=6)
    _centered(doc, PUBLISHER, size=12, color=TEAL, space_after=0)

    # --- Copyright page ---
    doc.add_page_break()
    year = _dt.date.today().year
    copyright_lines = [
        f"{BOOK_TITLE}: {BOOK_SUBTITLE}",
        "",
        f"Copyright \u00a9 {year} {PUBLISHER}. All rights reserved.",
        f"Copyright holder / context: {TBD}",
        "",
        f"ISBN (print): {TBD}",
        f"ISBN (e-book): {TBD}",
        f"Edition: {TBD}",
        "",
        "No part of this publication may be reproduced, distributed, or "
        "transmitted in any form or by any means without the prior written "
        "permission of the publisher.",
        "",
        f"Published by {PUBLISHER}.",
        "Manuscript compiled automatically from the repository sources via "
        "tools/build_manuscript.py.",
        f"Generated: {_dt.date.today().isoformat()}",
    ]
    for text in copyright_lines:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(text)
        run.font.size = Pt(9)
        if text == TBD or text.endswith(TBD):
            run.bold = True

    # --- Table of contents ---
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    note = doc.add_paragraph()
    run = note.add_run(
        "Final paginated table of contents to be supplied by DaScient, Inc. "
        f"{TBD}. In Microsoft Word, right-click the field below and choose "
        "\u201cUpdate Field\u201d to generate page numbers from the heading styles."
    )
    run.italic = True
    run.font.size = Pt(9)
    _insert_toc_field(doc)


def _insert_toc_field(doc: Document) -> None:
    """Insert a Word TOC field that Word can populate on demand."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "Right-click and choose \u201cUpdate Field\u201d to build the table of contents."
    )
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    for element in (fld_begin, instr, fld_sep, placeholder, fld_end):
        run._r.append(element)


def add_part_divider(doc: Document, part_name: str) -> None:
    doc.add_page_break()
    for _ in range(6):
        doc.add_paragraph()
    _centered(doc, part_name, size=26, bold=True, color=TEAL, space_after=0)


# --------------------------------------------------------------------------- #
# Document assembly
# --------------------------------------------------------------------------- #


def _configure_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)


def render_chapter(doc: Document, blocks: list[Block]) -> None:
    for block in blocks:
        if block.kind == "blank":
            continue
        if block.kind == "heading":
            _add_heading(doc, block)
        elif block.kind == "paragraph":
            para = doc.add_paragraph()
            add_inline(para, block.text)
        elif block.kind == "quote":
            para = doc.add_paragraph(style="Intense Quote")
            add_inline(para, block.text, base_italic=True)
        elif block.kind == "code":
            _add_code_block(doc, block)
        elif block.kind == "table":
            _add_table(doc, block)
        elif block.kind == "list":
            _add_list(doc, block)


def build(output: Path) -> tuple[Path, int]:
    chapters = collect_chapters()
    if not chapters:
        raise SystemExit("No chapters found in mkdocs.yml nav.")

    doc = Document()
    _configure_base_styles(doc)
    add_front_matter(doc)

    current_part: str | None = None
    for chapter in chapters:
        if chapter.part != current_part:
            add_part_divider(doc, chapter.part)
            current_part = chapter.part
        if not chapter.path.exists():
            raise SystemExit(f"Chapter source not found: {chapter.path}")
        blocks = parse_markdown(chapter.path.read_text(encoding="utf-8"))
        render_chapter(doc, blocks)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output, len(chapters)


def main() -> None:
    parser = argparse.ArgumentParser(
        description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=DEFAULT_OUTPUT,
        help=f"Destination .docx path (default: {DEFAULT_OUTPUT.relative_to(REPO_ROOT)})",
    )
    args = parser.parse_args()
    saved, count = build(args.output)
    print(f"Compiled {count} chapters into {saved}")


if __name__ == "__main__":
    main()
